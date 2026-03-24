"""Generate a professional project proposal Word document for the autonomous
AST data collection platform."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Project_Proposal_Autonomous_AST_Platform.docx")

# Paths to figures we want to embed
PLOTS_DIR = os.path.join(os.path.dirname(__file__), "results_crop_mlp", "plots")
PLOTS_NO_AMP_DIR = os.path.join(os.path.dirname(__file__), "results_crop_mlp", "plots_no_amp")

AGGREGATE_TRAJECTORY = os.path.join(PLOTS_DIR, "aggregate_timeseries_by_label.png")
NO_AMP_TRAJECTORY = os.path.join(PLOTS_NO_AMP_DIR, "aggregate_no_amp_trajectories.png")
FOLD0_TIMESERIES = os.path.join(PLOTS_DIR, "fold0_timeseries.png")


def set_cell_shading(cell, color_hex):
    """Set cell background colour."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F4E79")

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Styles --
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT PROPOSAL")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Autonomous Laboratory Platform for High-Throughput\n"
        "Generation of AI-Ready Antimicrobial Susceptibility\n"
        "Testing Data"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph("")
    doc.add_paragraph("")

    details = [
        ("Principal Investigator:", "Professor Till Bachmann"),
        ("Institution:", "Institute for Regeneration and Repair,\nUniversity of Edinburgh"),
        ("Project Originator\nand Industry Partner:", "Marcin Kedziera\nPhD Student, Biomedical AI CDT (2020 cohort)\nSchool of Informatics, University of Edinburgh\nFounder & Director, [Company Name]"),
        ("Co-Investigator:", "[Name \u2014 to be confirmed]"),
        ("Date:", "March 2026"),
        ("Version:", "1.0 \u2014 Pre-submission Draft"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_val = p.add_run(value)
        run_val.font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (placeholder)
    # =========================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Background and Motivation",
        "    2.1 The Antimicrobial Resistance Crisis",
        "    2.2 Current Limitations of AST",
        "    2.3 The Opportunity: Rapid Phenotypic AST via Microscopy and AI",
        "    2.4 The Data Bottleneck",
        "3. Proof of Concept: Existing Results",
        "    3.1 Experimental Setup",
        "    3.2 AI Pipeline",
        "    3.3 Key Results",
        "    3.4 No-Antibiotic Control Validation",
        "    3.5 Summary of Proof-of-Concept Evidence",
        "4. Project Vision and Objectives",
        "5. Technical Work Packages",
        "    5.1 WP1: Automated Bacterial Culture System",
        "    5.2 WP2: Next-Generation Imaging Hardware",
        "    5.3 WP3: Glass Microfluidic Chip with Reduced Channel Depth",
        "    5.4 WP4: Integrated Software and Data Pipeline",
        "    5.5 WP5: Dataset Generation at Scale",
        "    5.6 WP6: Validation and Quality Assurance",
        "6. Current Hardware Limitations and Proposed Improvements",
        "7. Dataset as a Commercial Product",
        "8. Impact and Market Opportunity",
        "9. Team and Expertise",
        "10. Budget Summary",
        "11. Timeline and Milestones",
        "12. Risk Register",
        "13. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "Antimicrobial resistance (AMR) is one of the most urgent threats to global public health. "
        "In 2019, bacterial AMR was associated with an estimated 4.95 million deaths worldwide, and "
        "projections suggest this toll will rise dramatically without intervention. A critical bottleneck "
        "in combating AMR is the speed of antimicrobial susceptibility testing (AST): current culture-based "
        "methods require 16\u201324 hours to determine which antibiotics will be effective against a patient\u2019s "
        "infection. During this delay, patients receive empirical therapy that may be inappropriate, "
        "contributing to treatment failures and further resistance selection."
    )

    doc.add_paragraph(
        "Rapid phenotypic AST \u2014 predicting antibiotic susceptibility within minutes rather than hours "
        "by directly observing bacterial morphological responses under a microscope \u2014 has the potential "
        "to transform clinical microbiology. However, the development of AI-driven rapid AST is critically "
        "constrained by the lack of large, standardised, high-quality training datasets. No such dataset "
        "currently exists at the scale required to train robust clinical-grade classifiers across multiple "
        "bacterial species and antibiotic combinations."
    )

    doc.add_paragraph(
        "This project proposes to build an autonomous laboratory platform for continuous, 24/7 generation "
        "of AI-ready microscopy data of bacteria under antibiotic exposure. The platform will automate the "
        "entire experimental pipeline \u2014 from bacterial culture preparation through microfluidic chip loading, "
        "time-lapse imaging, and data curation \u2014 and will address specific hardware limitations identified "
        "during proof-of-concept work. The resulting dataset will be the first of its kind: a large-scale, "
        "clinically validated, FAIR-compliant collection of phenotypic bacterial response data, suitable for "
        "training AI classifiers for rapid AST."
    )

    doc.add_paragraph(
        "Marcin Kedziera, a PhD student in the Biomedical AI CDT (2020 cohort) at the School of "
        "Informatics, University of Edinburgh, has already demonstrated that rapid phenotypic AST is "
        "feasible. Working under the supervision of Professor Till Bachmann at the Institute for "
        "Regeneration and Repair, Kedziera independently conceived the approach, designed and built "
        "the complete AI pipeline, generated all experimental data, and performed all analysis. "
        "Using a proof-of-concept system with significant hardware limitations \u2014 an 8-bit "
        "brightfield camera with an effective dynamic range of only 45 grey levels, a polymer "
        "microfluidic chip with a 20 \u00b5m channel depth that causes 30\u201340% of bacteria to be "
        "out of focus, and a fully manual workflow limited to approximately 3 experiments per "
        "working day \u2014 the AI pipeline achieves an AUROC of 0.76\u20130.80 on clinical E. coli "
        "samples exposed to ampicillin, with clear morphological signal emerging within 25\u201330 "
        "minutes. Crucially, a no-antibiotic control experiment confirms the classifier detects "
        "genuine drug-induced morphological changes (98.4% confidence) rather than strain identity. "
        "These results, obtained despite substantial hardware constraints, provide strong evidence "
        "that an improved platform generating data at scale will enable clinically relevant rapid "
        "AST across a broad range of pathogens and antibiotics."
    )

    p = doc.add_paragraph(
        "The dataset produced by this platform is the primary commercial output. It will be owned "
        "by a company founded by Kedziera, who holds the pre-existing intellectual property as the "
        "student inventor, and offered to major in vitro diagnostics (IVD) companies \u2014 including "
        "Siemens Healthineers, bioM\u00e9rieux, BD, and Thermo Fisher Scientific \u2014 who require "
        "high-quality training data to develop AI-driven rapid AST capabilities for their "
        "next-generation diagnostic instruments. These companies have the resources and expertise to "
        "develop diagnostic hardware, but they cannot build AI-driven AST without the training data "
        "this platform will produce. The UK government\u2019s AI for Science Strategy (November 2025) "
        "identifies autonomous laboratory platforms and high-value scientific datasets as national "
        "priorities, and has committed \u00a3137 million to AI research infrastructure including an "
        "open call on autonomous labs."
    )

    doc.add_page_break()

    # =========================================================================
    # 2. BACKGROUND AND MOTIVATION
    # =========================================================================
    doc.add_heading("2. Background and Motivation", level=1)

    doc.add_heading("2.1 The Antimicrobial Resistance Crisis", level=2)
    doc.add_paragraph(
        "Antimicrobial resistance occurs when bacteria evolve mechanisms to survive exposure to antibiotics "
        "that would normally kill them or inhibit their growth. The World Health Organization has declared "
        "AMR one of the top ten global public health threats facing humanity. In the UK alone, an estimated "
        "12,000 deaths per year are attributable to drug-resistant infections, and the O\u2019Neill Review "
        "projected that by 2050, AMR could cause 10 million deaths annually worldwide and cost the global "
        "economy up to $100 trillion."
    )
    doc.add_paragraph(
        "A key driver of resistance is inappropriate antibiotic prescribing. When clinicians prescribe "
        "broad-spectrum antibiotics because susceptibility results are not yet available, this creates "
        "selection pressure that promotes resistance. Faster, more accurate AST is therefore not merely a "
        "diagnostic improvement \u2014 it is a fundamental intervention against the spread of resistance."
    )

    doc.add_heading("2.2 Current Limitations of AST", level=2)
    doc.add_paragraph(
        "The gold-standard AST methods \u2014 broth microdilution and disk diffusion \u2014 are fundamentally "
        "growth-based: they require bacteria to be cultured in the presence of antibiotics for 16\u201324 "
        "hours, after which growth or inhibition is assessed visually or by optical density measurement. "
        "This approach has remained largely unchanged for decades. While automated systems (e.g., VITEK 2, "
        "BD Phoenix) have reduced hands-on time, the underlying requirement for bacterial growth means "
        "the minimum time to result remains many hours."
    )
    doc.add_paragraph(
        "Molecular methods (PCR-based resistance gene detection) offer speed but detect genotype, not "
        "phenotype. A bacterium carrying a resistance gene may not express it, and novel resistance "
        "mechanisms not covered by the gene panel will be missed. Phenotypic testing \u2014 directly observing "
        "how bacteria respond to antibiotic exposure \u2014 remains the most reliable approach, but it needs "
        "to be made faster."
    )

    doc.add_heading("2.3 The Opportunity: Rapid Phenotypic AST via Microscopy and AI", level=2)
    doc.add_paragraph(
        "When susceptible bacteria are exposed to an effective antibiotic, their morphology changes "
        "within minutes to hours: beta-lactam antibiotics (e.g., ampicillin) cause cell elongation, "
        "blebbing, and eventual lysis; fluoroquinolones cause nucleoid condensation and filamentation; "
        "aminoglycosides cause membrane disruption. These morphological changes are visible under a "
        "microscope well before the bacteria would have grown sufficiently to be detected by conventional "
        "growth-based AST."
    )
    doc.add_paragraph(
        "Recent advances in computer vision and deep learning make it feasible to detect these subtle "
        "morphological shifts automatically. Self-supervised learning methods such as DINO "
        "(self-distillation with no labels) can learn rich visual representations from unlabelled "
        "microscopy images, and Vision Transformers (ViTs) have been shown to outperform convolutional "
        "neural networks for bacterial classification from microscopy (Hallstrom et al., PLOS ONE 2025). "
        "The combination of time-lapse microscopy within microfluidic devices and AI-driven image analysis "
        "creates a viable path to rapid phenotypic AST."
    )

    doc.add_heading("2.4 The Data Bottleneck", level=2)
    doc.add_paragraph(
        "The critical limiting factor is not the AI methodology \u2014 it is the data. Training a robust "
        "clinical-grade AST classifier requires:"
    )
    bullets = [
        "Hundreds of bacterial strains across multiple species (E. coli, K. pneumoniae, "
        "S. aureus, P. aeruginosa, A. baumannii, and others on the WHO priority pathogen list)",
        "Multiple antibiotics per species at clinically relevant concentrations, including drugs "
        "from different mechanistic classes (beta-lactams, fluoroquinolones, aminoglycosides, "
        "carbapenems, polymyxins)",
        "Standardised imaging conditions to ensure reproducibility and minimise batch effects",
        "Clinical isolates with confirmed resistance profiles to provide ground truth labels",
        "Sufficient biological replicates per strain\u2013antibiotic combination to capture "
        "natural variation",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "No existing laboratory setup can generate data at this scale without full automation. "
        "Manual sample preparation, chip loading, and imaging is limited to approximately one experiment "
        "per hour. An autonomous platform operating 24/7 could increase throughput by an order of magnitude, "
        "generating the volume and diversity of data required to train classifiers that generalise across "
        "the clinical landscape."
    )

    doc.add_page_break()

    # =========================================================================
    # 3. PROOF OF CONCEPT
    # =========================================================================
    doc.add_heading("3. Proof of Concept: Existing Results", level=1)

    doc.add_paragraph(
        "A proof-of-concept system has been developed and validated, demonstrating that rapid phenotypic "
        "AST is achievable using time-lapse microscopy and deep learning, even with significant hardware "
        "limitations. This section presents the experimental setup, AI pipeline, and key results."
    )

    doc.add_heading("3.1 Experimental Setup", level=2)

    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Organism", "Escherichia coli (clinical isolates)"],
            ["Antibiotic", "Ampicillin, 16 mg/L"],
            ["Imaging modality", "Brightfield microscopy, 100x magnification"],
            ["Frame rate", "5 frames per second"],
            ["Image resolution", "1280 x 1024 pixels, grayscale"],
            ["Experiment duration", "1 hour per experiment"],
            ["Frames per experiment", "~14,500"],
            ["Total experiments", "42 (11 resistant, 16 susceptible, 15 test)"],
            ["Total bacterial strains", "15 unique strains (7 resistant, 8 susceptible)"],
            ["Total bacteria detected", "2.38 million individual crops"],
            ["Sample delivery", "Microfluidic channel with flowing media"],
        ],
        col_widths=[6, 10],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Bacteria are imaged within a microfluidic channel as they flow through the field of view "
        "under continuous media flow. Each bacterium is visible for approximately 2 seconds (~10 frames) "
        "before leaving the field of view. This setup prevents individual cell tracking but enables "
        "population-level analysis: at any given time point, hundreds of bacteria are visible, and the "
        "statistical properties of the population shift over time as susceptible bacteria respond to "
        "the antibiotic."
    )

    doc.add_heading("3.2 AI Pipeline", level=2)
    doc.add_paragraph(
        "The analysis pipeline operates in four stages:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Stage 1 \u2014 Bacteria Detection: ")
    run.bold = True
    p.add_run(
        "A finetuned YOLOv11 model with oriented bounding boxes detects bacteria in each frame "
        "and classifies them as focused, unfocused, or vertical. Only focused detections are retained. "
        "Each bacterium is cropped at its native pixel size onto a 128\u00d7128 canvas with reflected "
        "border padding to preserve size and shape information."
    )

    p = doc.add_paragraph()
    run = p.add_run("Stage 2 \u2014 Self-Supervised Feature Learning (DINO): ")
    run.bold = True
    p.add_run(
        "A Vision Transformer (ViT-Small, 384-dimensional embeddings) is pretrained using DINO "
        "self-supervised learning on ~207,000 bacteria crops. The model learns morphological "
        "representations without any labels by training a student network to match a momentum-averaged "
        "teacher across different augmented views of the same bacterium. Augmentations are calibrated to "
        "the data\u2019s narrow dynamic range (brightfield microscopy), with CLAHE contrast enhancement "
        "applied as a preprocessing step."
    )

    p = doc.add_paragraph()
    run = p.add_run("Stage 3 \u2014 Feature Extraction: ")
    run.bold = True
    p.add_run(
        "The pretrained backbone extracts a 384-dimensional feature vector for every detected bacterium "
        "across all 42 experiments, producing 2.38 million feature vectors."
    )

    p = doc.add_paragraph()
    run = p.add_run("Stage 4 \u2014 Classification: ")
    run.bold = True
    p.add_run(
        "Two classification approaches were evaluated. A per-crop MLP classifier predicts "
        "resistant/susceptible for each individual bacterium embedding, with experiment-level "
        "predictions derived from time-binned population voting. A population temporal classifier "
        "aggregates per-bin statistics (mean, standard deviation, skewness, kurtosis) through a "
        "transformer encoder with gated attention pooling."
    )

    doc.add_heading("3.3 Key Results", level=2)

    doc.add_paragraph(
        "All results use strain-holdout cross-validation (5 folds), where each fold holds out 2 "
        "resistant and 2 susceptible strains entirely. This ensures the classifier is evaluated on "
        "bacterial strains it has never seen, testing generalisation to novel clinical isolates."
    )

    # Results table
    doc.add_paragraph("")
    add_styled_table(doc,
        ["Classifier", "AUROC @ 60 min", "Key Feature"],
        [
            ["Population Temporal (Transformer + Stats)", "0.802 \u00b1 0.092", "Highest aggregate AUROC"],
            ["Per-Crop MLP", "0.764 \u00b1 0.141", "Clearest biological signal"],
            ["Sub-Sequence Sampling", "0.739 \u00b1 0.265", "Temporal signal emerges"],
            ["Delta Features", "0.744 \u00b1 0.081", "Lowest variance"],
            ["BiLSTM Temporal", "0.656 \u00b1 0.202", "\u2014"],
        ],
        col_widths=[6, 4, 6],
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Trajectory Analysis: ")
    run.bold = True
    p.add_run(
        "The most compelling evidence comes from trajectory analysis of the per-crop MLP classifier. "
        "Rather than evaluating a single prediction at the end of the experiment, we track the "
        "probability of resistance, P(R), over time for each experiment. The expected biological "
        "signal is that susceptible bacteria should initially appear resistant (the antibiotic has not "
        "yet taken effect) and progressively shift toward susceptible classification as morphological "
        "changes manifest."
    )

    doc.add_paragraph(
        "This is precisely what we observe:"
    )
    bullets = [
        "Susceptible experiments: Mean P(resistant) drops from 0.65 at t = 2 minutes to "
        "0.17 at t = 60 minutes",
        "Resistant experiments: Mean P(resistant) remains stable at ~0.70 throughout the "
        "full 60-minute observation period",
        "The susceptible and resistant trajectory curves cross the 0.5 decision boundary at "
        "approximately 25\u201330 minutes",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # Insert aggregate trajectory figure
    if os.path.exists(AGGREGATE_TRAJECTORY):
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(AGGREGATE_TRAJECTORY, width=Inches(5.0))
        caption = doc.add_paragraph(
            "Figure 1. Mean P(resistant) over time by true label, aggregated across all 5 strain-holdout "
            "folds. Shaded regions show \u00b11 standard deviation. Susceptible experiments (blue) show a "
            "clear downward trajectory as antibiotic-induced morphological changes accumulate, while "
            "resistant experiments (red) remain stable. Curves diverge from approximately 20 minutes."
        )
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(12)
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(9)

    # Insert per-fold figure
    if os.path.exists(FOLD0_TIMESERIES):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(FOLD0_TIMESERIES, width=Inches(5.0))
        caption = doc.add_paragraph(
            "Figure 2. Per-experiment P(resistant) trajectories for Fold 0 (holdout strains: EC58, EC87 "
            "resistant; EC36, EC39 susceptible). Solid lines represent true resistant experiments; dashed "
            "lines represent true susceptible experiments. Individual experiment trajectories confirm the "
            "population-level trend: susceptible experiments diverge downward over time."
        )
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(12)
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(9)

    doc.add_heading("3.4 No-Antibiotic Control Validation", level=2)
    doc.add_paragraph(
        "A critical question is whether the classifier learns genuine drug-induced morphological "
        "changes or simply memorises strain-specific visual fingerprints. To test this, five control "
        "experiments were conducted: susceptible strains (EC33, EC36, EC39) were imaged without any "
        "antibiotic exposure. These bacteria are genetically susceptible but have not been treated, "
        "so they should exhibit normal growth morphology \u2014 indistinguishable from resistant bacteria."
    )

    doc.add_paragraph(
        "The prediction logic is:"
    )
    bullets = [
        "If the classifier predicts untreated susceptible bacteria as resistant \u2192 it has "
        "learned to detect morphology (correct: untreated bacteria look healthy/resistant)",
        "If the classifier predicts them as susceptible \u2192 it has learned strain identity "
        "(incorrect: recognising the strain rather than the drug response)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph("")
    add_styled_table(doc,
        ["Classifier", "Mean P(R) @ 60 min", "% Classified Resistant", "Signal Source"],
        [
            ["Per-Crop MLP", "0.984 \u00b1 0.015", "100%", "Morphology"],
            ["BiLSTM Temporal", "0.785 \u00b1 0.361", "80%", "Morphology"],
            ["No-Count (morphology only)", "0.752 \u00b1 0.310", "72%", "Morphology"],
            ["Baseline (Transformer + Stats)", "0.610 \u00b1 0.402", "60%", "Mixed"],
        ],
        col_widths=[5, 3.5, 3.5, 3.5],
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        "The per-crop MLP classifier predicts all 5 untreated susceptible experiments as resistant "
        "with 98.4% confidence and negligible variance across all folds. This result holds regardless "
        "of whether the strain was in the training set or held out for that fold, providing strong "
        "evidence that the classifier has learned to detect the absence of drug-induced morphological "
        "damage rather than memorising strain appearance."
    )

    # Insert no-amp control figure
    if os.path.exists(NO_AMP_TRAJECTORY):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(NO_AMP_TRAJECTORY, width=Inches(5.0))
        caption = doc.add_paragraph(
            "Figure 3. No-antibiotic control: P(resistant) trajectories for susceptible strains imaged "
            "without ampicillin exposure, aggregated across folds. The per-crop MLP (top row) consistently "
            "classifies untreated bacteria as resistant with high confidence, confirming the classifier "
            "detects morphological changes induced by antibiotic exposure."
        )
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(12)
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(9)

    doc.add_heading("3.5 Summary of Proof-of-Concept Evidence", level=2)
    bullets = [
        "Rapid phenotypic AST for E. coli/ampicillin is achievable within 25\u201330 minutes using "
        "time-lapse microscopy and deep learning",
        "The AI classifier learns genuine drug-induced morphological changes, not strain identity "
        "(validated by no-antibiotic control experiments with 98.4% consistency)",
        "AUROC of 0.76\u20130.80 is achieved despite significant hardware limitations "
        "(low-contrast brightfield imaging, narrow dynamic range, limited field of view, manual "
        "sample preparation)",
        "The trajectory-based evaluation paradigm reveals biological signal that standard "
        "aggregate accuracy metrics obscure",
        "These results were obtained with only 42 experiments across 15 strains \u2014 "
        "performance is expected to improve substantially with more data",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    # =========================================================================
    # 4. PROJECT VISION AND OBJECTIVES
    # =========================================================================
    doc.add_heading("4. Project Vision and Objectives", level=1)

    doc.add_paragraph(
        "The vision of this project is to create the world\u2019s first autonomous laboratory platform "
        "purpose-built for continuous generation of AI-ready antimicrobial susceptibility testing data. "
        "This platform will operate 24/7, automating every step from bacterial culture preparation to "
        "curated dataset output, and will produce a large-scale, FAIR-compliant dataset that enables "
        "the development of clinically robust AI-driven rapid AST."
    )

    doc.add_paragraph("The specific objectives are:")

    objectives = [
        ("O1.", "Design and build an autonomous platform that automates bacterial culture, "
         "microfluidic chip loading, time-lapse microscopy, and data curation, enabling "
         "continuous 24/7 operation with minimal human intervention."),
        ("O2.", "Upgrade the imaging hardware to address limitations identified in the "
         "proof of concept: replace the current camera with a higher-resolution, higher-dynamic-range "
         "sensor; optimise the light path for phase-contrast or differential interference contrast "
         "(DIC) imaging; and improve the magnification optics for consistent focus across the "
         "field of view."),
        ("O3.", "Redesign the microfluidic chip to improve bacterial visibility, reduce flow "
         "artefacts, enable multi-channel parallel experiments, and support standardised loading "
         "protocols compatible with automation."),
        ("O4.", "Generate a dataset of at least 500 experiments across 50+ bacterial strains "
         "and 5+ antibiotic classes within the project period, with standardised imaging conditions "
         "and confirmed resistance profiles."),
        ("O5.", "Develop an integrated software pipeline for real-time quality control, automated "
         "data curation, and FAIR-compliant metadata generation, ensuring the dataset meets the "
         "requirements of downstream AI model development."),
        ("O6.", "Validate the platform\u2019s output by training improved AST classifiers on the "
         "generated dataset and demonstrating performance improvements over the proof-of-concept "
         "results."),
    ]
    for label, text in objectives:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(text)

    doc.add_page_break()

    # =========================================================================
    # 5. TECHNICAL WORK PACKAGES
    # =========================================================================
    doc.add_heading("5. Technical Work Packages", level=1)

    doc.add_heading("5.1 WP1: Automated Bacterial Culture System", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Objective: ")
    run.bold = True
    p.add_run(
        "Automate the preparation, incubation, and delivery of bacterial samples for imaging, "
        "enabling continuous unsupervised operation."
    )
    doc.add_paragraph(
        "The current proof-of-concept requires manual preparation of bacterial cultures, manual "
        "addition of antibiotics, manual loading of microfluidic chips, and manual initiation of "
        "each experiment. A researcher must be present throughout to monitor focus and handle any "
        "issues. This limits throughput to a maximum of approximately 3 experiments per working "
        "day \u2014 a fundamental bottleneck that prevents dataset generation at the scale required "
        "for clinical-grade AI training. WP1 will develop an automated liquid handling and culture "
        "system capable of:"
    )
    bullets = [
        "Automated inoculation from frozen stock into growth media",
        "Temperature-controlled incubation to target growth phase",
        "Automated antibiotic preparation at specified concentrations from stock solutions",
        "Automated mixing of bacterial suspension with antibiotic and loading into "
        "microfluidic chips",
        "Scheduling and queuing of multiple experiments for continuous operation",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "Full automation of the culture and sample preparation pipeline transforms the fundamental "
        "economics of data generation. A manual workflow producing 3 experiments per working day "
        "yields approximately 60 experiments per month. An automated platform running 24/7, even "
        "at a conservative 1 experiment per hour, produces over 700 experiments per month \u2014 "
        "more than a 10-fold increase. This difference is the distinction between a proof-of-concept "
        "dataset and a commercially valuable data asset."
    )
    doc.add_paragraph(
        "This work package requires expertise in laboratory automation, liquid handling robotics, "
        "and microbiological protocols. The system will be designed for biosafety Level 2 operation "
        "and will incorporate contamination detection and automated decontamination cycles."
    )

    doc.add_heading("5.2 WP2: Next-Generation Imaging Hardware", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Objective: ")
    run.bold = True
    p.add_run(
        "Replace the current imaging system with hardware optimised for high-contrast, "
        "high-resolution bacterial microscopy with consistent focus and expanded dynamic range."
    )
    doc.add_paragraph("The current hardware has several identified limitations:")

    add_styled_table(doc,
        ["Limitation", "Impact on Data Quality", "Proposed Solution"],
        [
            ["Low dynamic range camera\n(8-bit, narrow effective range\n~40\u201385 of 0\u2013255)",
             "Morphological detail compressed into ~45 grey levels;\nsubtle features such as cell wall texture,\nmembrane irregularities, and early blebbing\nare unresolvable; requires aggressive\npost-processing (CLAHE) to extract usable contrast",
             "Scientific-grade camera with 12\u201316-bit depth\nand higher quantum efficiency;\n1000\u201316000x more grey levels for\nmorphological discrimination"],
            ["Brightfield illumination only",
             "Low inherent contrast for unstained bacteria;\nphase information lost; nucleoid condensation\nand fine internal structure invisible",
             "Phase-contrast or DIC optics to enhance\nbacterial visibility without staining;\nenables detection of internal morphological\nchanges (nucleoid, membrane)"],
            ["Fixed focal plane\n+ 20 \u00b5m channel depth",
             "Bacteria at different depths in the 20 \u00b5m\nchannel fall in and out of focus;\n~30\u201340% of detected bacteria discarded\nas unfocused, reducing usable data",
             "Motorised Z-stage with autofocus;\ncombined with thin glass chip (WP3)\nto confine bacteria within focal plane"],
            ["Limited field of view",
             "Fewer bacteria visible per frame,\nreducing statistical power per time bin",
             "Larger sensor or automated stage\nscanning to capture wider area"],
            ["Manual focus adjustment",
             "Focus drift during 1-hour experiments\nrequires operator intervention;\nincompatible with 24/7 operation",
             "Hardware autofocus system with\nclosed-loop feedback for\nunattended operation"],
        ],
        col_widths=[4.5, 5, 5.5],
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        "The imaging system redesign will be informed by the specific requirements of the AI pipeline. "
        "Higher contrast and dynamic range directly improve the quality of DINO self-supervised features "
        "\u2014 the current system compresses all morphological information into approximately 45 grey "
        "levels, whereas a 12-bit camera provides 4,096 levels and a 16-bit camera provides 65,536. "
        "This expansion in dynamic range would make currently invisible features \u2014 bacterial cell "
        "wall texture, early-stage blebbing, membrane disruption, nucleoid condensation \u2014 directly "
        "resolvable. These are precisely the morphological changes induced by different antibiotic "
        "classes, and resolving them is expected to substantially improve classification accuracy "
        "and enable discrimination between different mechanisms of antibiotic action."
    )
    doc.add_paragraph(
        "Consistent focus, achieved through the combination of autofocus hardware and the thin glass "
        "microfluidic chip (WP3), eliminates the current reliance on the YOLO focused/unfocused classifier "
        "as a data quality filter."
    )

    doc.add_heading("5.3 WP3: Microfluidic Chip Redesign", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Objective: ")
    run.bold = True
    p.add_run(
        "Develop next-generation glass microfluidic chips with reduced channel depth, optimised "
        "for automated loading, maximal bacterial focus, and multi-experiment throughput."
    )
    doc.add_paragraph(
        "The current microfluidic chip uses a polymer-based single-channel configuration with a "
        "channel depth of approximately 20 \u00b5m. While functional for proof of concept, this depth "
        "is substantially greater than the focal plane of 100x microscope objectives (typically "
        "<1 \u00b5m depth of field). The consequence is severe: bacteria at different heights within "
        "the 20 \u00b5m channel fall in and out of focus, and in the current dataset approximately "
        "30\u201340% of all detected bacteria must be discarded as unfocused. This represents an "
        "enormous loss of usable data and introduces selection bias toward bacteria that happen to "
        "be near the focal plane at the moment of imaging."
    )

    p = doc.add_paragraph()
    run = p.add_run("Glass microfluidic chip with reduced channel depth. ")
    run.bold = True
    p.add_run(
        "The single most impactful hardware improvement is to transition to a glass microfluidic "
        "chip with a channel depth of 3\u20135 \u00b5m \u2014 thin enough to physically confine bacteria "
        "within or very close to the focal plane. At this depth, a rod-shaped E. coli bacterium "
        "(~1 \u00b5m diameter) has minimal room to move out of focus. Glass provides superior optical "
        "properties compared to polymer: lower autofluorescence, better refractive index matching "
        "with immersion oil objectives, and consistent optical flatness. Combined with the improved "
        "light path and higher-quality optics proposed in WP2, a thin glass channel would bring "
        "virtually all bacteria into sharp focus simultaneously, eliminating the need for the YOLO "
        "focused/unfocused classifier as a data quality gate and dramatically increasing the number "
        "of usable bacteria per frame."
    )

    p = doc.add_paragraph()
    run = p.add_run("Impact on morphological discrimination. ")
    run.bold = True
    p.add_run(
        "The combination of a shallower glass channel, improved phase-contrast or DIC optics, and "
        "a higher-bit-depth camera would unlock morphological features that are currently invisible "
        "or barely detectable. The proof-of-concept system operates with brightfield illumination "
        "and an 8-bit camera whose effective dynamic range spans only ~45 grey levels (40\u201385 "
        "of 0\u2013255). Under these conditions, subtle features such as bacterial cell wall texture, "
        "membrane irregularities, early-stage blebbing, and nucleoid condensation are compressed "
        "into a handful of pixel values and are largely unresolvable. A 12\u201316-bit camera with "
        "phase-contrast optics and a thin glass channel providing consistent focus would make these "
        "fine-grained morphological features directly visible, substantially improving the quality "
        "of DINO self-supervised feature representations and, consequently, downstream AST "
        "classification accuracy."
    )

    doc.add_paragraph("Additional chip improvements include:")
    bullets = [
        "Multi-channel chip design enabling parallel experiments (e.g., multiple antibiotics "
        "or concentrations simultaneously on a single chip)",
        "Optimised flow rate to increase bacterial residence time in the field of view "
        "(currently ~2 seconds), enabling more frames per bacterium and potentially supporting "
        "short-term individual tracking",
        "Standardised chip-to-microscope interface for automated loading and alignment, "
        "compatible with the robotic liquid handling system in WP1",
        "Integrated reference markers for automated focus calibration and drift correction",
        "Disposable design for single-use operation, preventing cross-contamination between "
        "experiments",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5.4 WP4: Integrated Software and Data Pipeline", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Objective: ")
    run.bold = True
    p.add_run(
        "Build the software infrastructure for autonomous operation, real-time quality control, "
        "and FAIR-compliant data output."
    )
    doc.add_paragraph(
        "The software pipeline will encompass:"
    )
    bullets = [
        "Experiment orchestration: scheduling, hardware control, and state management for "
        "continuous unattended operation",
        "Real-time quality monitoring: automated detection of focus drift, illumination "
        "changes, chip clogging, contamination, and other failure modes",
        "Automated data processing: on-the-fly bacteria detection, crop extraction, and "
        "feature extraction using the DINO pipeline",
        "FAIR-compliant metadata: automated recording of all experimental parameters "
        "(strain, antibiotic, concentration, imaging conditions, timestamps) in standardised "
        "formats",
        "Data storage and access: structured storage of raw images, processed crops, feature "
        "vectors, and metadata in a queryable database",
        "Dashboard and alerting: web-based monitoring interface for remote supervision and "
        "automated alerts for system failures",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5.5 WP5: Dataset Generation at Scale", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Objective: ")
    run.bold = True
    p.add_run(
        "Operate the platform to generate a large-scale, diverse dataset spanning multiple bacterial "
        "species, antibiotics, and resistance mechanisms."
    )
    doc.add_paragraph(
        "Once the platform is operational, dataset generation will proceed systematically across "
        "clinically relevant pathogen\u2013antibiotic combinations, prioritised by WHO critical and "
        "high priority pathogen lists. Initial targets include:"
    )

    add_styled_table(doc,
        ["Pathogen", "Antibiotic Classes", "Priority"],
        [
            ["Escherichia coli", "Beta-lactams, fluoroquinolones, aminoglycosides, carbapenems", "Critical"],
            ["Klebsiella pneumoniae", "Beta-lactams, carbapenems, polymyxins", "Critical"],
            ["Pseudomonas aeruginosa", "Carbapenems, fluoroquinolones, aminoglycosides", "Critical"],
            ["Staphylococcus aureus", "Methicillin (oxacillin), vancomycin, daptomycin", "High"],
            ["Acinetobacter baumannii", "Carbapenems, polymyxins, tigecycline", "Critical"],
        ],
        col_widths=[5, 7, 3],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Clinical isolates will be sourced through collaboration with NHS Lothian microbiology "
        "laboratories and the UK Health Security Agency (UKHSA), with confirmed resistance profiles "
        "from conventional AST as ground truth labels."
    )

    doc.add_heading("5.6 WP6: Validation and Quality Assurance", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Objective: ")
    run.bold = True
    p.add_run(
        "Validate the dataset quality through independent classifier training and benchmarking "
        "against proof-of-concept results."
    )
    doc.add_paragraph(
        "Validation will include: training improved AST classifiers on the expanded dataset and "
        "demonstrating performance improvement over the 42-experiment proof of concept; independent "
        "review of a random subset of experiments by clinical microbiologists; statistical analysis "
        "of batch effects, imaging consistency, and label accuracy; and publication of benchmark "
        "results to establish the dataset as a community resource."
    )

    doc.add_page_break()

    # =========================================================================
    # 6. CURRENT HARDWARE LIMITATIONS
    # =========================================================================
    doc.add_heading("6. Current Hardware Limitations and Proposed Improvements", level=1)

    doc.add_paragraph(
        "The proof-of-concept results were achieved despite substantial hardware limitations. "
        "These limitations constrain data quality and throughput but also demonstrate the robustness "
        "of the approach: if AUROC of 0.76\u20130.80 is achievable under these conditions, improved "
        "hardware should significantly enhance performance."
    )

    doc.add_paragraph(
        "The following table summarises the current limitations and the expected impact of "
        "proposed improvements:"
    )
    doc.add_paragraph("")

    add_styled_table(doc,
        ["Component", "Current State", "Limitation", "Proposed Improvement", "Expected Impact"],
        [
            ["Camera", "8-bit sensor,\neffective range\n40\u201385/255",
             "87% of dynamic\nrange unused;\nCLAHE needed;\ncell wall texture\ninvisible",
             "12\u201316-bit scientific\nCMOS sensor",
             "Resolve fine\nmorphological\nfeatures (cell wall\ntexture, membrane\nirregularities)"],
            ["Illumination", "Brightfield only",
             "Low inherent\nbacterial contrast;\nphase info lost",
             "Phase-contrast\nor DIC optics",
             "Direct morphological\nvisibility; nucleoid\ncondensation and\nblebbing resolvable"],
            ["Focus", "Manual, fixed\nfocal plane",
             "30\u201340% bacteria\ndiscarded as\nunfocused; drift\nover 1-hour expt",
             "Motorised Z-stage\nwith autofocus",
             "2\u20133x more usable\ndata per experiment;\nno operator needed"],
            ["Microfluidic\nchip", "Polymer, 20 \u00b5m\nchannel depth,\nsingle channel,\nmanual loading",
             "Channel 20x deeper\nthan focal plane;\n30\u201340% bacteria\nout of focus;\none experiment\nat a time",
             "Glass chip, 3\u20135 \u00b5m\nchannel depth;\nmulti-channel;\nautomated loading",
             "Near-100% bacteria\nin focus; parallel\nexperiments; better\noptical properties"],
            ["Sample\npreparation", "Fully manual:\nculture, antibiotic\naddition, chip\nloading",
             "Maximum ~3\nexperiments per\nworking day;\nresearcher must\nbe present",
             "Automated liquid\nhandling, culture\nincubation, and\nchip loading",
             "24/7 unattended\noperation; ~8x\nthroughput increase\n(24 vs 3 expts/day)"],
        ],
        col_widths=[2.5, 3, 3, 3.5, 3],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "The compound effect of these improvements is transformative for both data quantity and quality. "
        "The current manual workflow limits throughput to approximately 3 experiments per working day "
        "\u2014 a researcher must prepare cultures, load chips, initiate imaging, and monitor for focus "
        "drift throughout each 1-hour experiment. An automated platform running 24/7 could execute "
        "24 or more experiments per day, an 8-fold increase in throughput before accounting for the "
        "additional efficiency of multi-channel parallel chips."
    )
    doc.add_paragraph(
        "On the quality side, the transition from a 20 \u00b5m polymer channel to a 3\u20135 \u00b5m "
        "glass channel is the single most impactful change. By physically confining bacteria within "
        "the focal plane, it eliminates the largest source of data loss (30\u201340% of detections "
        "currently discarded as unfocused) and ensures every imaged bacterium is in sharp focus. "
        "Combined with phase-contrast or DIC optics and a high-bit-depth camera, this would enable "
        "the AI pipeline to detect subtle morphological features \u2014 cell wall texture, early "
        "blebbing, membrane irregularities, nucleoid condensation \u2014 that are currently invisible "
        "in the low-contrast, narrow-dynamic-range brightfield images. Given that the proof of "
        "concept already achieves AUROC 0.76\u20130.80 under these severe hardware constraints, the "
        "improvements described here should substantially enhance classification accuracy and reduce "
        "the time-to-prediction."
    )

    doc.add_page_break()

    # =========================================================================
    # 7. DATASET AS COMMERCIAL PRODUCT
    # =========================================================================
    doc.add_heading("7. Dataset as a Commercial Product", level=1)

    doc.add_paragraph(
        "The primary commercial output of this project is the dataset itself. Unlike a diagnostic "
        "device or a pharmaceutical compound, a large-scale, curated, clinically validated AI training "
        "dataset is a data asset that can be owned, licensed, and sold by the company that generates it."
    )

    doc.add_heading("7.1 Commercial Structure and Data Ownership", level=2)
    doc.add_paragraph(
        "The commercial exploitation of this project\u2019s output will be conducted through a "
        "company founded and directed by Marcin Kedziera, who originated the concept, developed the "
        "proof-of-concept platform, generated the initial dataset, and performed all data analysis "
        "demonstrating that rapid phenotypic AST is feasible. The pre-existing intellectual property "
        "\u2014 the proof-of-concept system, pipeline design, and associated know-how \u2014 was "
        "developed by Kedziera as a PhD student at the University of Edinburgh, and under the "
        "University\u2019s published IP policy, student-created IP belongs to the student by default."
    )
    doc.add_paragraph(
        "The proposed structure is as follows:"
    )
    bullets = [
        "The company will own all data generated by the platform. The dataset is the commercial "
        "product, and data ownership must reside with the company to enable licensing and sales "
        "to third-party customers.",
        "The University of Edinburgh will receive an equity stake in the company, consistent with "
        "the University Spinout Investment Terms (USIT) guide, in recognition of the use of "
        "university facilities and the academic contribution of the PI and co-investigators.",
        "The university retains rights to use all data and research outputs for academic, "
        "non-commercial research and publication purposes.",
        "The collaboration agreement \u2014 defining data ownership, equity allocation, commercial "
        "rights, and academic publication rights \u2014 will be negotiated and signed before the "
        "grant commences, not after.",
        "Kedziera, as founder and director of the company, will lead the technical delivery of "
        "the project and retain a majority shareholding commensurate with his role as originator "
        "of the idea, developer of the proof of concept, and ongoing technical lead.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "This structure ensures clean IP and data ownership that is attractive to future investors "
        "and acquirers. Large IVD companies will not purchase or license a dataset encumbered by "
        "complex university IP restrictions. A company that owns its data outright, with university "
        "equity as the mechanism for academic participation, is a significantly more investable and "
        "acquirable entity."
    )

    doc.add_heading("7.2 Why This Dataset Is Uniquely Valuable", level=2)

    bullets = [
        "No equivalent exists: There is no publicly available or commercially available large-scale "
        "microscopy dataset of bacteria under antibiotic exposure with confirmed resistance profiles "
        "and standardised imaging conditions.",
        "Difficult to replicate: Generating this dataset requires specialised hardware (automated "
        "microscopy platform), clinical microbiology expertise (to source and characterise strains), "
        "and sustained laboratory operation. Major diagnostic companies are not set up to do this "
        "foundational data generation internally.",
        "Essential for AI development: Every company developing AI-driven rapid AST \u2014 and there "
        "are many, given the market opportunity \u2014 needs training data of this type. The data is "
        "the fuel for their algorithms. These companies could undoubtedly build the diagnostic device "
        "themselves, but they cannot build it without training data.",
        "Growing in value: As additional strains, antibiotics, and resistance mechanisms are added, "
        "the dataset becomes more comprehensive and more valuable. The moat around this data asset "
        "deepens with every experiment.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("7.3 Target Customers", level=2)
    doc.add_paragraph(
        "The primary customers are major in vitro diagnostics (IVD) companies developing next-generation "
        "AST systems. These companies have the resources to build diagnostic hardware but lack the "
        "high-quality, large-scale phenotypic microscopy data required to train their AI classifiers:"
    )
    add_styled_table(doc,
        ["Company", "Relevant Products / Programmes", "Dataset Need"],
        [
            ["Siemens Healthineers", "Automated microbiology analysers; AI diagnostics R&D",
             "Training data for AI-driven AST modules"],
            ["bioM\u00e9rieux", "VITEK 2, BioFire; active AI/ML diagnostics programme",
             "Phenotypic data to complement molecular panels"],
            ["BD (Becton Dickinson)", "BD Phoenix; BD Kiestra automated microbiology",
             "Data for next-gen phenotypic AST algorithms"],
            ["Thermo Fisher Scientific", "Sensititre; clinical microbiology portfolio",
             "Standardised training data for AST classifier development"],
            ["Bruker", "MALDI-TOF MS (identification); expanding into AST",
             "Phenotypic data to build AST capability"],
            ["Accelerate Diagnostics", "Accelerate Pheno (morphokinetic AST)",
             "Expanded species/antibiotic coverage data"],
        ],
        col_widths=[4, 6, 5],
    )
    doc.add_paragraph("")

    doc.add_heading("7.4 Commercialisation Models", level=2)
    bullets = [
        "Exclusive licence: An exclusive licence to a specific IVD company for a defined field of use "
        "(e.g., automated clinical AST instruments), generating upfront payment plus royalties, "
        "while retaining rights for other applications and continued data generation.",
        "Non-exclusive licensing: Broad access to multiple companies at lower per-deal value but "
        "higher total revenue and no single-buyer dependency.",
        "Dataset-as-a-service: Subscription-based access to the dataset as it grows, suited to "
        "AI companies requiring continually updated training data as resistance patterns evolve.",
        "Acquisition: Outright sale of the dataset and platform IP to a major IVD company \u2014 "
        "the cleanest exit, potentially for millions of pounds, once the dataset reaches sufficient "
        "scale and species/antibiotic diversity.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    # =========================================================================
    # 8. IMPACT AND MARKET OPPORTUNITY
    # =========================================================================
    doc.add_heading("8. Impact and Market Opportunity", level=1)

    doc.add_heading("8.1 Clinical Impact", level=2)
    doc.add_paragraph(
        "Rapid phenotypic AST has the potential to reduce time-to-effective-therapy from 16\u201324 hours "
        "to under 1 hour. In sepsis, where each hour of delayed appropriate antibiotic therapy "
        "increases mortality by 7.6%, this represents a direct life-saving intervention. By enabling "
        "faster, targeted prescribing, rapid AST also reduces inappropriate broad-spectrum antibiotic "
        "use \u2014 a primary driver of AMR."
    )

    doc.add_heading("8.2 Market Opportunity", level=2)
    doc.add_paragraph(
        "The global AST market was valued at approximately $4.1 billion in 2023 and is projected to "
        "reach $6.5 billion by 2030, driven by increasing AMR prevalence and regulatory pressure for "
        "faster diagnostics. The AI in diagnostics segment is growing at over 20% CAGR. Companies "
        "investing in AI-driven AST face a common bottleneck: the lack of large, high-quality training "
        "datasets. This project directly addresses that bottleneck."
    )

    doc.add_heading("8.3 Strategic Alignment", level=2)
    doc.add_paragraph(
        "This project aligns with several major UK and international initiatives:"
    )
    bullets = [
        "UK AI for Science Strategy (November 2025): Identifies autonomous laboratory platforms "
        "and high-value scientific datasets as national priorities; \u00a3137M committed to AI "
        "research infrastructure including an open call on autonomous labs",
        "UK National Action Plan on AMR (2024\u20132029): Emphasises innovation in diagnostics "
        "and surveillance as pillars of the national AMR response",
        "WHO Global Action Plan on AMR: Calls for investment in rapid diagnostic tools to "
        "improve antibiotic stewardship globally",
        "UKRI FAIR Data Mandate: By 2030, all experimental data from UKRI facilities must meet "
        "FAIR principles \u2014 this platform is designed for FAIR-compliant data output from inception",
        "G7 AMR commitments: Continued funding and policy emphasis on novel diagnostic "
        "approaches for AMR",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    # =========================================================================
    # 9. TEAM AND EXPERTISE
    # =========================================================================
    doc.add_heading("9. Team and Expertise", level=1)

    doc.add_paragraph(
        "The project brings together interdisciplinary expertise spanning clinical microbiology, "
        "AMR diagnostics, AI/machine learning, optical engineering, and laboratory automation."
    )

    add_styled_table(doc,
        ["Role", "Name", "Affiliation", "Key Expertise"],
        [
            ["Project Originator,\nTechnical Lead,\nand Industry\nPartner", "Marcin Kedziera",
             "PhD Student, Biomedical\nAI CDT (2020 cohort),\nSchool of Informatics,\nUniversity of Edinburgh;\nFounder & Director,\n[Company Name]",
             "Originator of the entire project concept,\nstrategy, and proof of concept. Designed\nand built the AI pipeline (YOLO, DINO,\nViT, temporal classifiers), generated all\nexperimental data, performed all analysis.\nOwns pre-existing IP as student inventor."],
            ["Principal\nInvestigator", "Prof. Till Bachmann",
             "Institute for Regeneration\nand Repair, University\nof Edinburgh",
             "AMR diagnostics; biosensors; rapid point-of-\ncare testing. AMR Strategy Lead for Edinburgh\n"
             "Infectious Diseases; CARB-X Advisor; JPIAMR\nScientific Advisory Board Chair; Longitude\nPrize Judge."],
            ["Co-Investigator\n(AI/ML)", "[To be confirmed]",
             "School of Informatics,\nUniversity of Edinburgh",
             "Computer vision; deep learning; self-supervised\nlearning; medical image analysis."],
            ["Co-Investigator\n(Engineering)", "[To be confirmed]",
             "School of Engineering,\nUniversity of Edinburgh",
             "Optical systems; microscopy instrumentation;\nlaboratory automation; microfluidics."],
        ],
        col_widths=[2.5, 3, 4, 6],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Marcin Kedziera is the originator of this project in its entirety. As a PhD student in "
        "the Biomedical AI CDT (2020 cohort) at the School of Informatics, University of Edinburgh, "
        "he independently conceived the idea of using time-lapse microscopy with deep learning for "
        "rapid phenotypic AST, designed the experimental strategy, developed the complete AI pipeline "
        "(YOLO object detection, DINO self-supervised learning, Vision Transformer feature extraction, "
        "population temporal classification, and trajectory-based evaluation), generated all "
        "experimental data, and performed all data analysis that demonstrates the feasibility of "
        "rapid phenotypic AST. The proof-of-concept results presented in Section 3 are entirely "
        "his work. Under the University of Edinburgh\u2019s published IP policy, intellectual property "
        "created by students belongs to the student by default. The pre-existing IP \u2014 the "
        "proof-of-concept platform, pipeline architecture, and methodology \u2014 therefore belongs "
        "to Kedziera and will be held by his company."
    )
    doc.add_paragraph(
        "Professor Bachmann\u2019s position as AMR Strategy Lead for Edinburgh Infectious Diseases, "
        "his advisory role with CARB-X, and his chairmanship of the JPIAMR Scientific Advisory Board "
        "provide unparalleled access to clinical isolate collections, international AMR networks, and "
        "diagnostic industry contacts. His existing relationships with EPSRC, BBSRC, Wellcome Trust, "
        "and PACE funding bodies provide strategic advantage in grant applications. His role as PI "
        "provides the academic leadership required by UKRI funding bodies, while Kedziera\u2019s "
        "company serves as the industry partner responsible for commercial exploitation of the "
        "dataset."
    )

    doc.add_page_break()

    # =========================================================================
    # 10. BUDGET SUMMARY
    # =========================================================================
    doc.add_heading("10. Budget Summary", level=1)

    doc.add_paragraph(
        "The following is an indicative budget for a 3-year project. Actual figures will be refined "
        "during detailed application preparation and will depend on the specific funding scheme."
    )

    doc.add_paragraph("")
    add_styled_table(doc,
        ["Category", "Year 1", "Year 2", "Year 3", "Total"],
        [
            ["Staff costs", "", "", "", ""],
            ["  Research Associate (AI/ML)", "\u00a342,000", "\u00a343,000", "\u00a344,000", "\u00a3129,000"],
            ["  Research Associate (Engineering)", "\u00a342,000", "\u00a343,000", "\u00a344,000", "\u00a3129,000"],
            ["  Research Technician", "\u00a332,000", "\u00a333,000", "\u00a334,000", "\u00a399,000"],
            ["Equipment", "", "", "", ""],
            ["  Scientific camera + optics", "\u00a365,000", "\u2014", "\u2014", "\u00a365,000"],
            ["  Motorised microscope stage", "\u00a325,000", "\u2014", "\u2014", "\u00a325,000"],
            ["  Liquid handling robot", "\u00a380,000", "\u2014", "\u2014", "\u00a380,000"],
            ["  Compute (GPU server)", "\u00a320,000", "\u2014", "\u2014", "\u00a320,000"],
            ["Consumables", "", "", "", ""],
            ["  Microfluidic chips", "\u00a38,000", "\u00a315,000", "\u00a320,000", "\u00a343,000"],
            ["  Culture media & antibiotics", "\u00a35,000", "\u00a310,000", "\u00a315,000", "\u00a330,000"],
            ["  Bacterial strain acquisition", "\u00a310,000", "\u00a38,000", "\u00a35,000", "\u00a323,000"],
            ["Travel & dissemination", "\u00a35,000", "\u00a37,000", "\u00a38,000", "\u00a320,000"],
            ["Indirect costs (estimated)", "\u00a390,000", "\u00a375,000", "\u00a370,000", "\u00a3235,000"],
            ["", "", "", "", ""],
            ["TOTAL (indicative)", "\u00a3324,000", "\u00a3334,000", "\u00a3240,000", "\u00a3898,000"],
        ],
        col_widths=[6, 2.5, 2.5, 2.5, 2.5],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: These figures are indicative and will be adjusted based on the specific requirements "
        "and eligible cost categories of the target funding scheme. Indirect costs are estimated at "
        "the University of Edinburgh\u2019s standard FEC rate. Equipment costs reflect market estimates "
        "for research-grade components."
    )

    doc.add_page_break()

    # =========================================================================
    # 11. TIMELINE AND MILESTONES
    # =========================================================================
    doc.add_heading("11. Timeline and Milestones", level=1)

    add_styled_table(doc,
        ["Phase", "Period", "Milestones", "Deliverables"],
        [
            ["Phase 1:\nDesign &\nProcurement",
             "Months\n1\u20136",
             "Camera and optics specified and ordered;\nmicrofluidic chip v2 designed;\nautomation system architecture finalised",
             "Hardware specification document;\nchip design files;\nsystem architecture document"],
            ["Phase 2:\nBuild &\nIntegrate",
             "Months\n7\u201312",
             "Imaging system assembled and calibrated;\nliquid handling integrated;\nfirst automated experiment completed",
             "Assembled platform;\ncalibration report;\nfirst automated dataset"],
            ["Phase 3:\nCommission\n& Optimise",
             "Months\n13\u201318",
             "Platform running 24/7 with E. coli/ampicillin;\ndata pipeline operational;\nquality metrics established",
             "50+ experiments generated;\nQC dashboard;\ndata pipeline documentation"],
            ["Phase 4:\nScale to\nMulti-Species",
             "Months\n19\u201330",
             "Additional species and antibiotics integrated;\n200+ experiments generated;\nfirst classifier benchmark",
             "Multi-species dataset;\nbenchmark results;\ninterim publication"],
            ["Phase 5:\nValidation\n& Release",
             "Months\n31\u201336",
             "500+ experiments completed;\nindependent validation;\ndataset publication;\ncommercial discussions",
             "Full dataset;\nvalidation report;\npublication(s);\nbusiness case"],
        ],
        col_widths=[3, 2, 5.5, 4.5],
    )

    doc.add_page_break()

    # =========================================================================
    # 12. RISK REGISTER
    # =========================================================================
    doc.add_heading("12. Risk Register", level=1)

    add_styled_table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["Hardware integration delays",
             "Medium", "Medium",
             "Modular design allowing incremental integration;\ncontinue proof-of-concept data collection\nin parallel"],
            ["Imaging quality insufficient\nfor new species",
             "Low", "High",
             "Iterative optimisation; phase-contrast/DIC provides\nfundamentally higher contrast than brightfield;\nfallback to CLAHE-enhanced brightfield if needed"],
            ["Microfluidic chip clogging\nor contamination",
             "Medium", "Medium",
             "Disposable chips; automated flushing and\ndecontamination protocols; redundant channels"],
            ["Classifier performance does\nnot improve with more data",
             "Low", "High",
             "Proof of concept already shows 0.76\u20130.80 AUROC\nwith 42 experiments; learning curves suggest\nperformance is data-limited, not method-limited"],
            ["Clinical isolate availability",
             "Low", "Medium",
             "NHS Lothian partnership; UKHSA strain collections;\nProf. Bachmann\u2019s international AMR network"],
            ["IP or commercialisation\ndisputes",
             "Medium", "High",
             "Pre-negotiated collaboration agreement before\ngrant start; independent IP legal advice;\nclear separation of pre-existing and new IP"],
            ["Key personnel departure",
             "Low", "Medium",
             "Documented protocols and codebase;\nknowledge transfer plan;\ncross-training within team"],
        ],
        col_widths=[4, 2, 2, 7],
    )

    doc.add_page_break()

    # =========================================================================
    # 13. REFERENCES
    # =========================================================================
    doc.add_heading("13. References", level=1)

    refs = [
        "Murray, C.J. et al. (2022). Global burden of bacterial antimicrobial resistance in 2019: "
        "a systematic analysis. The Lancet, 399(10325), 629\u2013655.",
        "O\u2019Neill, J. (2016). Tackling Drug-Resistant Infections Globally: Final Report and "
        "Recommendations. Review on Antimicrobial Resistance.",
        "Lim, S.H. et al. (2018). Phenotypic Antimicrobial Susceptibility Testing with Deep "
        "Learning Video Microscopy. Analytical Chemistry, 90(22), 13485\u201313490.",
        "Zagajewski, A. et al. (2023). Deep learning and single-cell phenotyping for rapid "
        "antimicrobial susceptibility detection. Communications Biology, 6, 1235.",
        "Hallstrom, S. et al. (2025). Rapid label-free identification of bacterial species using "
        "microfluidic time-lapse microscopy. PLOS ONE, 20(1), e0316592.",
        "Caron, M. et al. (2021). Emerging Properties in Self-Supervised Vision Transformers "
        "(DINO). Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV).",
        "Oquab, M. et al. (2024). DINOv2: Learning Robust Visual Features without Supervision. "
        "Transactions on Machine Learning Research.",
        "Cell-DINO (2025). Cell-DINO: Self-supervised pretraining for cell fluorescent microscopy "
        "image analysis. PLOS Computational Biology.",
        "Ilse, M. et al. (2018). Attention-based Deep Multiple Instance Learning. Proceedings of "
        "the 35th International Conference on Machine Learning (ICML).",
        "Guo, C. et al. (2017). On Calibration of Modern Neural Networks. Proceedings of the 34th "
        "International Conference on Machine Learning (ICML).",
        "UK Government (2025). AI for Science: A Strategy for UK Research and Innovation. Department "
        "for Science, Innovation and Technology.",
        "Kumar, A. et al. (2006). Duration of hypotension before initiation of effective "
        "antimicrobial therapy is the critical determinant of survival in human septic shock. "
        "Critical Care Medicine, 34(6), 1589\u20131596.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"[{i}] ")
        run_num.bold = True
        run_num.font.size = Pt(9)
        run_text = p.add_run(ref)
        run_text.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(4)

    # =========================================================================
    # Save
    # =========================================================================
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
